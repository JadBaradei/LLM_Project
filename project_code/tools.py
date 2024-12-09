import os
from langchain.tools import tool  # Decorator to create tools
from langchain_google_genai import GoogleGenerativeAIEmbeddings  # Google's embeddings
from langchain_chroma import Chroma  # Vector database
from langchain_community.document_loaders import PyPDFLoader  # Loads and parses PDF files
from get_key import get_api_key
from docx import Document as DocxDocument
from langchain.schema import Document
from serpapi import GoogleSearch
import re
import pandas as pd

plot_type = None
plotted_figure = None

# This function works as a search vector depending on what the user uploads. It has a cache mechanism as well where it will only scan the newly uploaded files 
# not all files everytime.
@tool
def search_vector_db(query: str) -> str:
    """
    Search the vector database for documents similar to the query.
    Args:
        query (str): The search query string to find relevant documents
    Returns:
        str: A concatenated string of the top 5 most similar document contents found in the vector database
    """
    
    # Variable holding the List[Document] to add to the vectordb
    added_documents = []

    # Initialize embedding model
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=get_api_key())
    
    # Initialize/connect to vector database
    vector_store = Chroma(
        collection_name = "embeddings",
        embedding_function = embeddings,
        persist_directory = "./vector_db",
        collection_metadata = {"hnsw:space": "cosine"}  # Use cosine similarity
    )

    directory_path = os.path.join(os.path.dirname(__file__), '../', 'db')
    directory_details_path = os.path.join(directory_path, 'db_details')

    # Load last scan timestamps from a file
    last_scan_file = os.path.join(directory_details_path, 'last_scan.txt')
    if os.path.exists(last_scan_file):
        with open(last_scan_file, 'r') as f:
            last_scan_times = eval(f.read())  # Read and parse stored times
    else:
        last_scan_times = {}

    current_scan_times = {}

    # Loop through all files in the directory
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if os.path.isfile(file_path):
            mod_time = os.path.getmtime(file_path)
            current_scan_times[filename] = mod_time
            
            # Check if the file was added/modified after the last scan
            if filename not in last_scan_times or last_scan_times[filename] < mod_time:
                #Debug print
                print(f"I am scanning the file: {filename}")
                if filename.endswith('.pdf'):
                    loader = PyPDFLoader(file_path)
                    documents = loader.load()
                    added_documents.extend(documents)
                elif filename.endswith('.docx'):
                    docx_file = DocxDocument(file_path)
                    documents = [
                        Document(
                            page_content=para.text,
                            metadata={"source": file_path}
                        )
                        for para in docx_file.paragraphs if para.text.strip()
                    ]
                    added_documents.extend(documents)
                elif filename.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as file:
                        content = file.read()
                    documents = [
                        Document(
                            page_content=content,
                            metadata={"source": file_path}
                        )
                    ]
                    added_documents.extend(documents)

    # Update last scan times
    with open(last_scan_file, 'w') as f:
        f.write(str(current_scan_times))

    # Add the cached documents to the vector store with their embeddings
    if(len(added_documents) > 0):
        vector_store.add_documents(added_documents)

    # Debug print
    print("Searching the vector database for: ", query)
    
    # Perform similarity search and get top 5 results
    result = vector_store.similarity_search(query=query, k=5)
    result_str = "\n".join([doc.page_content for doc in result])
    
    return result_str

@tool
def plot_excel_sheet(query: str) -> str:
    """
    Draws a plot given the name of an uploaded excel sheet and displays it to the user.

    Args:
        query (str): The name of the excel sheet to plot.
    Returns:
        str: A confirmation message of whether the plot was successfull or not.
    """
    excel_path = search_db_by_name(query)
    print("Searched db for: ", excel_path)

    try:
        df = pd.read_excel(excel_path)
    except:
        return "Invalid excel file"

    try:
        if plot_type == "bar":
            ax = df.plot.bar()
        elif plot_type == "line":
            ax = df.plot.line()
        elif plot_type == "scatter":
            ax = df.plot.scatter(x=df.columns[0], y=df.columns[1])
        else:
            raise ValueError("Invalid plot type specified.")
    except:
        return "Could not draw plot"
    
    global plotted_figure
    plotted_figure = ax.get_figure()
    return f"Successfully drew plot of type {plot_type} as you selected"


@tool
def search_google_scholar(query: str) -> str:
    """
    Searches Google Scholar for research papers related to the user's query, formats the results and provides a brief summary of each paper found.
    
    Args:
        query (str): The search query string provided by the user.
    
    Returns:
        str: A formatted string containing research paper titles, snippets, and links.
    """
    # Debug print
    print("Searching google scholar for: ", query)

    parsed_results = google_scholar_query(query)

    if not parsed_results:
        return "No research results were found for your query."

    # Format the results
    formatted_results = []
    for idx, item in enumerate(parsed_results, start=1):
        title = item.get("title", "No title available")
        snippet = item.get("snippet", "No snippet available")
        link = item.get("link", "No link available")
        formatted_results.append(
            f"Result {idx}:\n"
            f"Title: {title}\n"
            f"Briefing: {snippet}\n"
            f"Link: {link}\n"
        )

    return "\n".join(formatted_results)

# The role of this function is to save the added files in the 'added' folder1
def save_uploaded_files(uploaded_files) -> str:
    
    # The db folder
    directory_path = os.path.join(os.path.dirname(__file__), '../', 'db')
    # The caching folder
    db_details_path = os.path.join(directory_path, 'db_details')
    # Check if the folder exists, create it if not
    if not os.path.exists(directory_path):
        os.makedirs(directory_path)
        os.makedirs(db_details_path)

    saved_files = []  # List to store paths of saved files

    # Loop through the uploaded files and save them
    for uploaded_file in uploaded_files:
        # Get the file name
        file_name = uploaded_file.name
        # Create the full file path
        file_path = os.path.join(directory_path, file_name)
        
        # Check if the file already exists in the directory
        if os.path.exists(file_path):
            saved_files.append(f"- File '{file_name}' skipped, it already exists")
            continue  # Skip this file if it already exists

        # Write the file to the specified folder
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())  # Save file content

        # Add the file path to the list of saved files
        saved_files.append(f"- {file_name}")

    return "\n".join(saved_files)

def google_scholar_query(query,api_key="9b1904997abe4f6501079460c47f8159d6ea5365ef04b3e476a56d7bd5da888c"):
    params = {
    "engine": "google_scholar",
    "q": query,
    "api_key": api_key
    }
    search = GoogleSearch(params)
    results = search.get_dict()
    organic_results = results.get("organic_results", [])
    parsed_results = [
        {
            "title": item.get("title", "No title available"),
            "link": item.get("link", "No link available"),
            "snippet": item.get("snippet", "No snippet available")
        }
        for item in organic_results
    ]
    if not parsed_results:
        return "Error: the api key must have been depleted, or no search results were found"
    return parsed_results

def search_db_by_name(name):

    directory_path = os.path.join(os.path.dirname(__file__), '../', 'db')

    for filename in os.listdir(directory_path):
        if contains_same_characters(filename, name):
            if not filename.endswith(".xlsx"):
                return None
            else:
                return os.path.join(directory_path, filename)

def contains_same_characters(str1, str2):
    # Remove any non-alphanumeric characters from str1
    cleaned_str1 = re.sub(r'[^a-zA-Z0-9]', '', str1)
    # Convert both strings to lowercase for case-insensitive comparison 
    cleaned_str1 = cleaned_str1.lower()
    cleaned_str2 = str2.lower()

    print(cleaned_str2 in cleaned_str1)
    # Check if the cleaned str2 exists within cleaned str1
    return cleaned_str2 in cleaned_str1

def set_plot_type(plt_type):
    global plot_type
    plot_type = plt_type

def get_plotted_figure():
    global plotted_figure
    return plotted_figure
